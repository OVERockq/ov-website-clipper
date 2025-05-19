"""
Web to Ebook Converter
Version: 0.0.4
Description: 웹 페이지를 다양한 형식(EPUB, Markdown, DOC, PDF)의 전자책으로 변환하는 도구
"""

import os
import time
import zipfile
from typing import List, Dict, Union
import requests
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from ebooklib import epub
from deep_translator import GoogleTranslator
from tqdm import tqdm
import logging
from urllib.parse import urljoin
from docx import Document
from docx.shared import Inches
import markdown
from weasyprint import HTML
import tempfile
import re

class WebToEbook:
    def __init__(self, title: str, base_url: str, content_selector: str, menu_selector: str = None, txt_content: str = None, translator_type: str = 'browser', papago_id: str = None, papago_secret: str = None, openai_key: str = None, deepl_key: str = None, font_family: str = 'Noto Sans KR'):
        self.title = title
        self.base_url = base_url
        self.content_selector = content_selector
        self.menu_selector = menu_selector
        self.txt_content = txt_content
        self.translator_type = translator_type
        self.papago_id = papago_id
        self.papago_secret = papago_secret
        self.openai_key = openai_key
        self.deepl_key = deepl_key
        self.font_family = font_family
        self.pages = []
        self.setup_logging()
        self.setup_driver()

    def setup_logging(self):
        logging.basicConfig(
            level=logging.DEBUG,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)

    def setup_driver(self):
        chrome_options = Options()
        chrome_options.add_argument('--headless')
        chrome_options.add_argument('--no-sandbox')
        chrome_options.add_argument('--disable-dev-shm-usage')
        chrome_options.add_argument('--disable-gpu')
        chrome_options.add_argument('--disable-extensions')
        chrome_options.add_argument('--disable-software-rasterizer')
        
        service = Service('/usr/local/bin/chromedriver')
        self.driver = webdriver.Chrome(service=service, options=chrome_options)

    def get_menu_links(self) -> List[str]:
        """메뉴에서 링크를 추출합니다."""
        if not self.menu_selector:
            return []
            
        try:
            self.driver.get(self.base_url)
            time.sleep(2)
            
            if self.menu_selector.startswith('//'):
                menu = self.driver.find_element("xpath", self.menu_selector)
            else:
                menu = self.driver.find_element("css selector", self.menu_selector)
                
            links = menu.find_elements("tag name", "a")
            return [link.get_attribute('href') for link in links if link.get_attribute('href')]
            except Exception as e:
            self.logger.error(f"메뉴 링크 추출 중 오류 발생: {str(e)}")
            return []
        
    def get_page_content(self, url: str = None) -> str:
        """페이지의 본문 내용을 추출합니다."""
        if url:
        self.driver.get(url)
        else:
            self.driver.get(self.base_url)
        time.sleep(2)
        
        try:
            if self.content_selector.startswith('//'):
            content = self.driver.find_element("xpath", self.content_selector)
            else:
                content = self.driver.find_element("css selector", self.content_selector)
                
            content_html = content.get_attribute('innerHTML')
            self.logger.debug(f"페이지 내용 추출 - URL: {url or self.base_url}")
            return content_html
        except Exception as e:
            self.logger.error(f"페이지 내용 추출 중 오류 발생: {str(e)}")
            return ""

    def process_txt_content(self) -> List[str]:
        """텍스트 파일 내용을 페이지로 분리합니다."""
        if not self.txt_content:
            return []
            
        return [page.strip() for page in self.txt_content.split('\n\n') if page.strip()]

    def create_index_page(self) -> str:
        """목차 페이지를 생성합니다."""
        index_html = f"""
        <h1>목차</h1>
        <ul>
        """
        
        for i, page in enumerate(self.pages, 1):
            title = page.get('title', f'페이지 {i}')
            index_html += f'<li><a href="#page{i}">{title}</a></li>\n'
            
        index_html += "</ul>"
        return index_html

    def create_epub(self, content: str, output_path: str):
        """EPUB 형식의 전자책을 생성합니다."""
        book = epub.EpubBook()
        
        # 메타데이터 설정
        book.set_identifier('id123456')
        book.set_title(self.title)
        book.set_language('ko')
        
        # CSS 스타일 추가
        style = f'''
        body {{ 
            font-family: "{self.font_family}", sans-serif;
            line-height: 1.8;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 2em;
        }}
        img, video {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 1.5em auto;
            border-radius: 4px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        pre, code {{
            max-width: 100%;
            overflow-x: auto;
            word-break: break-all;
            white-space: pre-wrap;
            box-sizing: border-box;
        }}
        h1 {{ 
            font-size: 2em;
            margin: 1.5em 0 1em;
            color: #1a1a1a;
            border-bottom: 1px solid #eaecef;
            padding-bottom: 0.3em;
        }}
        h2 {{ 
            font-size: 1.5em;
            margin: 1.2em 0 0.8em;
            color: #1a1a1a;
        }}
        h3 {{ 
            font-size: 1.25em;
            margin: 1em 0 0.6em;
            color: #1a1a1a;
        }}
        p {{
            margin: 0.8em 0;
            line-height: 1.8;
        }}
        pre {{
            background-color: #f6f8fa;
            border-radius: 6px;
            padding: 16px;
            overflow: auto;
            font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, monospace;
            font-size: 0.9em;
            line-height: 1.45;
            margin: 1em 0;
        }}
        code {{
            font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, monospace;
            background-color: rgba(27, 31, 35, 0.05);
            padding: 0.2em 0.4em;
            border-radius: 3px;
            font-size: 0.9em;
        }}
        pre code {{
            background-color: transparent;
            padding: 0;
            border-radius: 0;
        }}
        blockquote {{
            margin: 1em 0;
            padding: 0 1em;
            color: #6a737d;
            border-left: 0.25em solid #dfe2e5;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
        }}
        th, td {{
            border: 1px solid #dfe2e5;
            padding: 6px 13px;
        }}
        th {{
            background-color: #f6f8fa;
        }}
        hr {{
            height: 0.25em;
            padding: 0;
            margin: 24px 0;
            background-color: #e1e4e8;
            border: 0;
        }}
        '''
        css = epub.EpubItem(uid="style_default", file_name="style/default.css", media_type="text/css", content=style)
        book.add_item(css)
        
        # 표지 페이지 생성
        cover = epub.EpubHtml(title='표지', file_name='cover.xhtml', lang='ko')
        cover.content = f'''
        <html>
        <head>
            <link rel="stylesheet" href="style/default.css" type="text/css" />
        </head>
        <body>
            <h1>{self.title}</h1>
        </body>
        </html>
        '''
        book.add_item(cover)
        
        # 목차 페이지 생성
        if len(self.pages) > 1:
            index = epub.EpubHtml(title='목차', file_name='index.xhtml', lang='ko')
            index.content = f'''
            <html>
            <head>
                <link rel="stylesheet" href="style/default.css" type="text/css" />
            </head>
            <body>
                {self.create_index_page()}
            </body>
            </html>
            '''
            book.add_item(index)
        
        # 각 페이지 처리
        chapters = []
        for i, page in enumerate(self.pages, 1):
            chapter = epub.EpubHtml(
                title=page.get('title', f'페이지 {i}'),
                file_name=f'page{i}.xhtml',
                lang='ko'
            )
            
            # 이미지 처리
            soup = BeautifulSoup(self.clean_for_reading_mode(page['content']), 'html.parser')
            for img in soup.find_all('img'):
                img_url = urljoin(self.base_url, img.get('src', ''))
                try:
                    response = requests.get(img_url)
                    if response.status_code == 200:
                        img_filename = f"images/{hash(img_url)}.png"
                        img_item = epub.EpubItem(
                            uid=f"img_{hash(img_url)}",
                            file_name=img_filename,
                            media_type="image/png",
                            content=response.content
                        )
                        book.add_item(img_item)
                        img['src'] = img_filename
                        
                        # 이미지 캡션 추가
                        if img.get('alt'):
                            caption = soup.new_tag('p')
                            caption['class'] = 'image-caption'
                            caption.string = img['alt']
                            img.insert_after(caption)
                except Exception as e:
                    self.logger.error(f"이미지 처리 중 오류 발생: {str(e)}")
            
            # 코드 블록 처리
            for pre in soup.find_all('pre'):
                if pre.find('code'):
                    code = pre.find('code')
                    # 언어 지정이 있는 경우
                    if code.get('class'):
                        lang = code['class'][0].replace('language-', '')
                        pre['data-lang'] = lang
            
            chapter.content = f'''
            <html>
            <head>
                <link rel="stylesheet" href="style/default.css" type="text/css" />
            </head>
            <body>
                {str(soup)}
            </body>
            </html>
            '''
            book.add_item(chapter)
            chapters.append(chapter)
        
        # 목차 생성
        book.toc = [cover]
        if len(self.pages) > 1:
            book.toc.append(index)
        book.toc.extend(chapters)
        
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # 스파인 정의
        book.spine = ['nav', cover]
        if len(self.pages) > 1:
            book.spine.append(index)
        book.spine.extend(chapters)
        
        # EPUB 파일 생성
        epub.write_epub(output_path, book, {})

    def create_markdown(self, content: str, output_path: str):
        """Markdown 형식의 문서를 생성합니다."""
        if len(self.pages) > 1:
            # ZIP 파일로 묶기
            with zipfile.ZipFile(output_path, 'w') as zipf:
                # 목차 생성
                index_content = "# 목차\n\n"
                for i, page in enumerate(self.pages, 1):
                    title = page.get('title', f'페이지 {i}')
                    index_content += f"{i}. [{title}]({i:02d}_{title}.md)\n"
                
                # 목차 파일 추가
                zipf.writestr("00_index.md", index_content)
                
                # 각 페이지 처리
                for i, page in enumerate(self.pages, 1):
                    title = page.get('title', f'페이지 {i}')
                    filename = f"{i:02d}_{title}.md"
                    
                    soup = BeautifulSoup(self.clean_for_reading_mode(page['content']), 'html.parser')
                    md_content = ""
                    
                    # 모든 요소를 순서대로 처리
                    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'pre', 'blockquote', 'img', 'table']):
                        if element.name.startswith('h'):
                            level = int(element.name[1])
                            md_content += f"{'#' * level} {element.get_text().strip()}\n\n"
                        elif element.name == 'p':
                            text = element.get_text().strip()
                            if text:
                                md_content += f"{text}\n\n"
                        elif element.name == 'pre':
                            code = element.find('code')
                            if code:
                                lang = code.get('class', [''])[0].replace('language-', '') if code.get('class') else ''
                                md_content += f"```{lang}\n{code.get_text()}\n```\n\n"
                            else:
                                md_content += f"```\n{element.get_text()}\n```\n\n"
                        elif element.name == 'blockquote':
                            md_content += f"> {element.get_text().strip()}\n\n"
                        elif element.name == 'img':
                            alt = element.get('alt', '')
                            src = urljoin(self.base_url, element.get('src', ''))
                            md_content += f"![{alt}]({src})\n\n"
                            if alt:
                                md_content += f"*{alt}*\n\n"
                        elif element.name == 'table':
                            md_content += self._convert_table_to_markdown(element)
                    
                    zipf.writestr(filename, md_content)
        else:
            # 단일 페이지 처리
            soup = BeautifulSoup(self.clean_for_reading_mode(content), 'html.parser')
        with open(output_path, 'w', encoding='utf-8') as f:
                for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'pre', 'blockquote', 'img', 'table']):
                    if element.name.startswith('h'):
                        level = int(element.name[1])
                        f.write(f"{'#' * level} {element.get_text().strip()}\n\n")
                    elif element.name == 'p':
                        text = element.get_text().strip()
                        if text:
                            f.write(f"{text}\n\n")
                    elif element.name == 'pre':
                        code = element.find('code')
                        if code:
                            lang = code.get('class', [''])[0].replace('language-', '') if code.get('class') else ''
                            f.write(f"```{lang}\n{code.get_text()}\n```\n\n")
                        else:
                            f.write(f"```\n{element.get_text()}\n```\n\n")
                    elif element.name == 'blockquote':
                        f.write(f"> {element.get_text().strip()}\n\n")
                    elif element.name == 'img':
                        alt = element.get('alt', '')
                        src = urljoin(self.base_url, element.get('src', ''))
                        f.write(f"![{alt}]({src})\n\n")
                        if alt:
                            f.write(f"*{alt}*\n\n")
                    elif element.name == 'table':
                        f.write(self._convert_table_to_markdown(element))

    def _convert_table_to_markdown(self, table):
        """HTML 테이블을 Markdown 형식으로 변환합니다."""
        md = []
        rows = table.find_all('tr')
        
        # 헤더 처리
        headers = []
        for th in rows[0].find_all(['th', 'td']):
            headers.append(th.get_text().strip())
        md.append('| ' + ' | '.join(headers) + ' |')
        md.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
        
        # 데이터 행 처리
        for row in rows[1:]:
            cells = []
            for cell in row.find_all(['td', 'th']):
                cells.append(cell.get_text().strip())
            md.append('| ' + ' | '.join(cells) + ' |')
        
        return '\n'.join(md) + '\n\n'

    def create_doc(self, content: str, output_path: str):
        """Word 문서 형식으로 저장합니다."""
        doc = Document()
        
        # 기본 스타일 설정
        styles = doc.styles
        styles['Normal'].font.name = self.font_family
        styles['Normal'].font.size = Pt(11)
        styles['Normal'].paragraph_format.line_spacing = 1.5
        
        # 제목 스타일 설정
        for i in range(1, 7):
            style_name = f'Heading {i}'
            if style_name in styles:
                styles[style_name].font.name = self.font_family
                styles[style_name].font.size = Pt(16 - i)
                styles[style_name].font.bold = True
        
        # 코드 블록 스타일 설정
        if 'Code' not in styles:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = 'Consolas'
            code_style.font.size = Pt(10)
            code_style.paragraph_format.space_before = Pt(12)
            code_style.paragraph_format.space_after = Pt(12)
            code_style.paragraph_format.left_indent = Pt(24)
            code_style.paragraph_format.right_indent = Pt(24)
            code_style.paragraph_format.background_color = RGBColor(246, 248, 250)
        
        soup = BeautifulSoup(self.clean_for_reading_mode(content), 'html.parser')
        
        # 모든 요소를 순서대로 처리
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'pre', 'blockquote', 'img', 'table']):
            if element.name.startswith('h'):
                # 제목 태그 처리
                level = int(element.name[1])
                doc.add_heading(element.get_text().strip(), level=level)
            elif element.name == 'p':
                # 단락 처리
                text = element.get_text().strip()
                if text:  # 빈 단락은 건너뛰기
                    p = doc.add_paragraph(text)
                    p.style = 'Normal'
            elif element.name == 'pre':
                # 코드 블록 처리
                code = element.find('code')
                if code:
                    lang = code.get('class', [''])[0].replace('language-', '') if code.get('class') else ''
                    p = doc.add_paragraph(code.get_text(), style='Code')
                    if lang:
                        p.add_run(f'\n// {lang}').italic = True
                else:
                    doc.add_paragraph(element.get_text(), style='Code')
            elif element.name == 'blockquote':
                # 인용구 처리
                p = doc.add_paragraph(element.get_text().strip())
                p.style = 'Quote'
            elif element.name == 'img':
            # 이미지 처리
                try:
                    img_url = urljoin(self.base_url, element.get('src', ''))
                    response = requests.get(img_url)
                    if response.status_code == 200:
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                            tmp.write(response.content)
                            doc.add_picture(tmp.name, width=Inches(6))
                            if element.get('alt'):
                                p = doc.add_paragraph(element['alt'])
                                p.style = 'Caption'
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    self.logger.error(f"이미지 처리 중 오류 발생: {str(e)}")
            elif element.name == 'table':
                # 테이블 처리
                table = doc.add_table(rows=0, cols=0)
                table.style = 'Table Grid'
                
                # 헤더 처리
                header_row = table.add_row()
                for th in element.find_all('th'):
                    cell = header_row.add_cell()
                    cell.text = th.get_text().strip()
                    cell.paragraphs[0].style = 'Table Heading'
                
                # 데이터 행 처리
                for tr in element.find_all('tr')[1:]:
                    row = table.add_row()
                    for td in tr.find_all('td'):
                        cell = row.add_cell()
                        cell.text = td.get_text().strip()
        
        doc.save(output_path)

    def create_pdf(self, content: str, output_path: str):
        """PDF 형식의 문서를 생성합니다."""
        soup = BeautifulSoup(self.clean_for_reading_mode(content), 'html.parser')
        
        # 이미지 처리
        for img in soup.find_all('img'):
            try:
                img_url = urljoin(self.base_url, img.get('src', ''))
                response = requests.get(img_url)
                if response.status_code == 200:
                    # 임시 파일로 이미지 저장
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                        tmp.write(response.content)
                        # 이미지 태그의 src를 로컬 파일 경로로 변경
                        img['src'] = f"file://{tmp.name}"
            except Exception as e:
                self.logger.error(f"이미지 처리 중 오류 발생: {str(e)}")
                # 이미지 로드 실패 시 alt 텍스트로 대체
                img.replace_with(soup.new_string(f"[이미지: {img.get('alt', '')}]"))
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@400;700&display=swap');
                @page {{
                    margin: 2.5cm;
                    @top-right {{
                        content: counter(page);
                    }}
                }}
                body {{ 
                    font-family: '{self.font_family}', sans-serif;
                    line-height: 1.8;
                    color: #333;
                    max-width: 800px;
                    margin: 0 auto;
                }}
                h1 {{ 
                    font-size: 2em;
                    margin: 1.5em 0 1em;
                    color: #1a1a1a;
                    border-bottom: 1px solid #eaecef;
                    padding-bottom: 0.3em;
                }}
                h2 {{ 
                    font-size: 1.5em;
                    margin: 1.2em 0 0.8em;
                    color: #1a1a1a;
                }}
                h3 {{ 
                    font-size: 1.25em;
                    margin: 1em 0 0.6em;
                    color: #1a1a1a;
                }}
                p {{
                    margin: 0.8em 0;
                    line-height: 1.8;
                }}
                pre {{
                    background-color: #f6f8fa;
                    border-radius: 6px;
                    padding: 16px;
                    overflow: auto;
                    font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, monospace;
                    font-size: 0.9em;
                    line-height: 1.45;
                    margin: 1em 0;
                }}
                code {{
                    font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, monospace;
                    background-color: rgba(27, 31, 35, 0.05);
                    padding: 0.2em 0.4em;
                    border-radius: 3px;
                    font-size: 0.9em;
                }}
                pre code {{
                    background-color: transparent;
                    padding: 0;
                    border-radius: 0;
                }}
                img {{
                    max-width: 100%;
                    height: auto;
                    display: block;
                    margin: 1.5em auto;
                    border-radius: 4px;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                }}
                .image-caption {{
                    text-align: center;
                    color: #666;
                    font-style: italic;
                    margin-top: 0.5em;
                }}
                blockquote {{
                    margin: 1em 0;
                    padding: 0 1em;
                    color: #6a737d;
                    border-left: 0.25em solid #dfe2e5;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 1em 0;
                }}
                th, td {{
                    border: 1px solid #dfe2e5;
                    padding: 6px 13px;
                }}
                th {{
                    background-color: #f6f8fa;
                }}
                hr {{
                    height: 0.25em;
                    padding: 0;
                    margin: 24px 0;
                    background-color: #e1e4e8;
                    border: 0;
                }}
            </style>
        </head>
        <body>
            {str(soup)}
        </body>
        </html>
        """
        
        try:
            # PDF 생성
            HTML(string=html_content).write_pdf(
                output_path,
                presentational_hints=True,
                optimize_size=('fonts', 'images'),
                zoom=1
            )
            self.logger.info(f"PDF 파일이 생성되었습니다: {output_path}")
        except Exception as e:
            self.logger.error(f"PDF 생성 중 오류 발생: {str(e)}")
            raise

    def translate_text(self, text: str, target_lang: str) -> str:
        """텍스트를 지정된 언어로 번역합니다."""
        try:
            translator = GoogleTranslator(source='auto', target=target_lang)
            return translator.translate(text)
        except Exception as e:
            self.logger.error(f"번역 중 오류 발생: {str(e)}")
            return text

    def translate_html_text(self, html, target_lang, translator_type):
        soup = BeautifulSoup(html, 'html.parser')
        for elem in soup.find_all(text=True):
            parent = elem.parent.name
            if parent not in ['code', 'pre', 'script', 'style', 'table']:
                text = elem.strip()
                if text:
                    try:
                        if translator_type == 'papago':
                            translated = self.translate_with_papago(text, target_lang)
                        elif translator_type == 'gpt':
                            translated = self.translate_with_gpt(text, target_lang)
                        elif translator_type == 'deepl':
                            translated = self.translate_with_deepl(text, target_lang)
                        else:
                            translated = GoogleTranslator(source='auto', target=target_lang).translate(text)
                        elem.replace_with(translated)
                    except Exception as e:
                        self.logger.error(f'번역 오류: {e}')
        return str(soup)

    def translate_with_papago(self, text, target_lang):
        url = "https://openapi.naver.com/v1/papago/n2mt"
        headers = {
            "X-Naver-Client-Id": self.papago_id,
            "X-Naver-Client-Secret": self.papago_secret
        }
        data = {
            "source": "ko" if target_lang != "ko" else "en",
            "target": target_lang,
            "text": text
        }
        response = requests.post(url, headers=headers, data=data)
        if response.status_code == 200:
            return response.json()['message']['result']['translatedText']
        else:
            return text

    def translate_with_gpt(self, text, target_lang):
        import openai
        openai.api_key = self.openai_key
        prompt = f"Translate the following text to {target_lang}:\n{text}"
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}]
            )
            return response.choices[0].message['content'].strip()
        except Exception as e:
            self.logger.error(f'GPT 번역 오류: {e}')
            return text

    def translate_with_deepl(self, text, target_lang):
        url = "https://api-free.deepl.com/v2/translate"
        headers = {"Authorization": f"DeepL-Auth-Key {self.deepl_key}"}
        data = {"text": text, "target_lang": target_lang.upper()}
        response = requests.post(url, headers=headers, data=data)
        if response.status_code == 200:
            return response.json()['translations'][0]['text']
        else:
            return text

    def clean_for_reading_mode(self, html):
        soup = BeautifulSoup(html, 'html.parser')

        # SVG, 아이콘, 불필요한 장식 제거 및 대체
        for svg in soup.find_all('svg'):
            svg.decompose()
        for icon in soup.find_all(['i', 'span']):
            if 'icon' in icon.get('class', []) or 'material-icons' in icon.get('class', []):
                alt = icon.get('aria-label') or icon.get('title') or '[ICON]'
                icon.replace_with(soup.new_string(alt))
        # 불필요한 nav, aside, footer, button, form, script, style, header, noscript, iframe, input, select, textarea, label, ul.menu, div.menu 등 제거
        for tag in soup.find_all(['nav', 'aside', 'footer', 'button', 'form', 'script', 'style', 'header', 'noscript', 'iframe', 'input', 'select', 'textarea', 'label']):
            tag.decompose()
        for tag in soup.find_all(class_=['menu', 'sidebar', 'navigation', 'breadcrumb', 'toc', 'toolbar', 'pagination', 'logo', 'search', 'ads', 'ad', 'banner', 'sponsor']):
            tag.decompose()
        # 코드블록 언어 표시
        for pre in soup.find_all('pre'):
            code = pre.find('code')
            if code and code.get('class'):
                lang = code['class'][0].replace('language-', '')
                pre['data-lang'] = lang
        # 이미지 alt 캡션 추가
        for img in soup.find_all('img'):
            if img.get('alt'):
                caption = soup.new_tag('p')
                caption['class'] = 'image-caption'
                caption.string = img['alt']
                img.insert_after(caption)
        # 본문, 제목, 표, 코드, 이미지, 인용구만 남기고 나머지 제거
        allowed = ['h1','h2','h3','h4','h5','h6','p','pre','code','img','blockquote','ul','ol','li','table','thead','tbody','tr','th','td','hr']
        for tag in soup.find_all(True):
            if tag.name not in allowed:
                tag.unwrap()
        # [SVG] 문자열 후처리로 모두 삭제
        result = str(soup)
        result = result.replace('[SVG]', '')
        return result

    def process(self, output_path: str, target_lang: str = None, output_format: str = 'epub'):
        """전체 프로세스를 실행합니다."""
        try:
            self.logger.info("페이지 내용 수집 중...")
            
            # 텍스트 파일 내용 처리
            if self.txt_content:
                pages = self.process_txt_content()
                for page in pages:
                    self.pages.append({
                        'title': '페이지',
                        'content': f'<div>{page}</div>'
                    })
            else:
                # 메뉴 선택자가 있는 경우 메뉴 링크만 처리
                if self.menu_selector:
                    menu_links = self.get_menu_links()
                    for link in menu_links:
                        page_content = self.get_page_content(link)
                        if page_content:
                            self.pages.append({
                                'title': link.split('/')[-1],
                                'content': page_content
                            })
                else:
                    # 메뉴 선택자가 없는 경우 메인 페이지만 처리
                    main_content = self.get_page_content()
                    self.pages.append({
                        'title': self.title,
                        'content': main_content
                    })
            
            # 번역 처리
            if target_lang and self.pages:
                for i, page in enumerate(self.pages):
                    page['content'] = self.translate_html_text(page['content'], target_lang, self.translator_type)
            
            self.logger.info(f"{output_format.upper()} 파일 생성 중...")
            
            # 출력 형식에 따라 적절한 변환 함수 호출
            if output_format.lower() == 'epub':
                self.create_epub(self.pages[0]['content'], output_path)
            elif output_format.lower() == 'markdown':
                self.create_markdown(self.pages[0]['content'], output_path)
            elif output_format.lower() == 'doc':
                self.create_doc(self.pages[0]['content'], output_path)
            elif output_format.lower() == 'pdf':
                self.create_pdf(self.pages[0]['content'], output_path)
            else:
                raise ValueError(f"지원하지 않는 출력 형식입니다: {output_format}")
            
            self.logger.info(f"문서가 생성되었습니다: {output_path}")
            
        except Exception as e:
            self.logger.error(f"처리 중 오류 발생: {str(e)}")
        finally:
            self.driver.quit()

def main():
    # 사용 예시
    converter = WebToEbook(
        title="Cursor Documentation",
        base_url="https://docs.cursor.com/",
        content_selector="main",
        menu_selector="nav",  # 메뉴 선택자 추가
        translator_type="browser"
    )
    
    converter.process(
        output_path="cursor_docs.pdf",
        target_lang="en",
        output_format="pdf"
    )

if __name__ == "__main__":
    main() 