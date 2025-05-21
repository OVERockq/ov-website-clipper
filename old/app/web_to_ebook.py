"""
Web to Ebook Converter
Version: 0.0.4
Description: 웹 페이지를 다양한 형식(EPUB, Markdown, DOC, PDF)의 전자책으로 변환하는 도구
"""

import os
import time
import zipfile
from typing import List, Dict, Union
import requests
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from ebooklib import epub
from deep_translator import GoogleTranslator
from tqdm import tqdm
import logging
from urllib.parse import urljoin
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from weasyprint import HTML
import tempfile
import re

class WebToEbook:
    def __init__(self, title: str, base_url: str, content_selector: str, menu_selector: str = None,  
                 translator_type: str = 'none', target_lang: str = None, # target_lang 추가
                 papago_id: str = None, papago_secret: str = None,  
                 openai_key: str = None, deepl_key: str = None, font_family: str = 'Noto Sans KR'):
        self.title = title
        self.base_url = base_url
        self.content_selector = content_selector
        self.menu_selector = menu_selector
        self.translator_type = translator_type
        self.target_lang = target_lang # target_lang 저장
        self.papago_id = papago_id # papago_id 저장 추가
        self.papago_secret = papago_secret
        self.openai_key = openai_key
        self.deepl_key = deepl_key
        self.font_family = font_family
        self.pages = []
        self.setup_logging()
        
        self.setup_driver()

    def setup_logging(self):
        logging.basicConfig(
            level=logging.DEBUG,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)

    def setup_driver(self):
        chrome_options = Options()
        chrome_options.add_argument('--headless')
        chrome_options.add_argument('--no-sandbox')
        chrome_options.add_argument('--disable-dev-shm-usage')
        chrome_options.add_argument('--disable-extensions')
        
        # 'translator_type'이 'browser'일 때 Chrome 내장 번역 기능 활성화
        if self.translator_type == 'browser' and self.target_lang:
            self.logger.info(f"Chrome 내장 번역을 위해 브라우저 언어를 {self.target_lang}로 설정하고 번역 기능을 활성화합니다.")
            chrome_options.add_argument(f"--lang={self.target_lang}")
            # 번역을 활성화하고, 특정 언어 페이지를 항상 대상 언어로 번역하도록 설정 시도
            # 원본 언어를 알 수 없으므로, Chrome이 자동 감지 후 target_lang으로 번역하도록 유도
            # 번역 관련 prefs 설정 강화
            prefs = {
                "profile.default_content_setting_values.popups": 0,
                "profile.default_content_settings.translate": 1, # 1: 허용, 2: 차단
                "translate_whitelists": {}, # 초기화 후 동적으로 설정
                "translate":{"enabled":"true"}
            }

            # 대상 언어가 'ko'일 때, 영어, 일본어, 중국어, 미정의 언어를 한국어로 번역하도록 설정
            if self.target_lang == 'ko':
                prefs["translate_whitelists"] = {
                    "en": "ko", "ja": "ko", "zh-CN": "ko", "zh-TW": "ko", "und": "ko"
                }
            # 대상 언어가 'en'일 때, 한국어, 일본어, 중국어, 미정의 언어를 영어로 번역하도록 설정
            elif self.target_lang == 'en':
                 prefs["translate_whitelists"] = {
                    "ko": "en", "ja": "en", "zh-CN": "en", "zh-TW": "en", "und": "en"
                }
            # 기타 대상 언어의 경우, 영어와 미정의 언어를 해당 대상 언어로 번역하도록 기본 설정
            else:
                prefs["translate_whitelists"] = {
                    "en": self.target_lang, "und": self.target_lang
                }
            
            # 만약 whitelist의 키와 값이 같다면 (예: "en":"en"), 해당 항목 제거
            keys_to_delete = [key for key, value in prefs["translate_whitelists"].items() if key == value]
            for key in keys_to_delete:
                del prefs["translate_whitelists"][key]

            self.logger.info(f"Chrome prefs 설정: {prefs}")
            chrome_options.add_experimental_option("prefs", prefs)
            # 일부 성공 사례에서 언급된 추가 옵션 (효과 미지수)
            chrome_options.add_experimental_option('excludeSwitches', ['enable-automation'])
            chrome_options.add_argument('--disable-blink-features=AutomationControlled')

        service = Service('/usr/local/bin/chromedriver') # 경로가 시스템에 맞게 설정되어 있는지 확인
        self.driver = webdriver.Chrome(service=service, options=chrome_options)

    def get_menu_links(self) -> List[str]:
        """메뉴에서 링크를 추출합니다."""
        if not self.menu_selector:
            return []
        try:
            self.driver.get(self.base_url)
            time.sleep(2)
            if self.menu_selector.startswith('//'):
                menu = self.driver.find_element("xpath", self.menu_selector)
            else:
                menu = self.driver.find_element("css selector", self.menu_selector)
            links = menu.find_elements("tag name", "a")
            return [link.get_attribute('href') for link in links if link.get_attribute('href')]
        except Exception as e:
            self.logger.error(f"메뉴 링크 추출 중 오류 발생: {str(e)}")
            return []
        
    def get_page_content(self, url: str = None) -> str:
        """페이지의 본문 내용을 추출합니다."""
        effective_url = url or self.base_url
        self.driver.get(effective_url)

        # 'translator_type'이 'browser'일 때 번역 적용 대기
        if self.translator_type == 'browser' and self.target_lang:
            self.logger.debug(f"내장 브라우저 번역 시도: URL='{effective_url}', 대상 언어='{self.target_lang}'")
            self.logger.info(f"Chrome 내장 번역 ({self.target_lang}) 적용 대기 시작...")
            
            # 번역 완료를 기다리는 더 확실한 방법 (예: html 태그의 lang 속성 변경 감지)
            from selenium.webdriver.support.ui import WebDriverWait
            from selenium.webdriver.support import expected_conditions as EC
            from selenium.webdriver.common.by import By
            
            original_lang = self.driver.find_element(By.TAG_NAME, "html").get_attribute("lang")
            self.logger.debug(f"원본 페이지 HTML lang 속성: {original_lang}")

            # 최대 15초까지 html의 lang 속성이 target_lang으로 변경되기를 기다림
            # 또는 lang 속성이 비어있지 않고, original_lang과 다른 경우도 번역된 것으로 간주 (일부 사이트는 target_lang과 정확히 일치 안 할 수 있음)
            try:
                # WebDriverWait 조건: html lang 속성 변경 또는 body에 'translated-' 클래스 존재
                WebDriverWait(self.driver, 15).until(
                    lambda driver: (
                        driver.find_element(By.TAG_NAME, "html").get_attribute("lang") == self.target_lang or
                        (driver.find_element(By.TAG_NAME, "html").get_attribute("lang") != "" and
                         driver.find_element(By.TAG_NAME, "html").get_attribute("lang") != original_lang and
                         self.target_lang != original_lang) or
                        ("translated-" in driver.find_element(By.TAG_NAME, "body").get_attribute("class"))
                    )
                )
                final_html_lang = self.driver.find_element(By.TAG_NAME, "html").get_attribute("lang")
                body_class = self.driver.find_element(By.TAG_NAME, "body").get_attribute("class")
                self.logger.info(f"Chrome 내장 번역 적용 완료 또는 변경 감지. HTML lang: {final_html_lang}, Body class: {body_class}")
                # WebDriverWait가 성공했으므로, 추가적인 긴 대기는 불필요할 수 있음. DOM 안정화를 위해 짧은 시간만 유지하거나 제거.
                time.sleep(0.5) # DOM 안정화를 위한 매우 짧은 추가 대기
            except Exception as e_wait:
                self.logger.warning(f"Chrome 내장 번역 완료 감지 시간 초과 또는 실패: {e_wait}. 현재 페이지 내용으로 계속 진행합니다. (추가 대기 3초)")
                time.sleep(3) # WebDriverWait 실패 시, 번역이 느리게 적용될 가능성을 고려한 최소한의 대기 (기존 5초에서 줄임)
        else:
            time.sleep(2) 
        
        try:
            if self.content_selector == "__AUTO_SINGLE_PAGE__":
                # 단일 페이지 자동 인식 모드: body 전체를 가져오고 clean_for_reading_mode에 의존
                self.logger.info(f"단일 페이지 자동 인식 모드로 body 전체 내용 추출 - URL: {effective_url}")
                content_html = self.driver.find_element("tag name", "body").get_attribute('innerHTML')
            elif self.content_selector.startswith('//'): # XPath
                content = self.driver.find_element("xpath", self.content_selector)
                content_html = content.get_attribute('innerHTML')
            else: # CSS Selector
                content = self.driver.find_element("css selector", self.content_selector)
                content_html = content.get_attribute('innerHTML')
            
            self.logger.debug(f"페이지 내용 추출 완료 - URL: {effective_url}")
            return content_html
        except Exception as e:
            self.logger.error(f"페이지 내용 추출 중 오류 발생: {str(e)}")
            return ""

    def create_index_page(self) -> str:
        """목차 페이지를 생성합니다."""
        index_html = f"""
        <h1>목차</h1>
        <ul>
        """
        
        for i, page in enumerate(self.pages, 1):
            title = page.get('title', f'페이지 {i}')
            index_html += f'<li><a href="#page{i}">{title}</a></li>\n'
            
        index_html += "</ul>"
        return index_html

    def create_epub(self, content: str, output_path: str):
        """EPUB 형식의 전자책을 생성합니다."""
        book = epub.EpubBook()
        
        # 메타데이터 설정
        book.set_identifier('id123456')
        book.set_title(self.title)
        book.set_language('ko')
        
        # URL에서 도메인 추출하여 저자로 설정
        try:
            from urllib.parse import urlparse
            parsed_url = urlparse(self.base_url)
            domain = parsed_url.netloc
            if domain:
                book.add_author(domain)
            else:
                book.add_author("Unknown") # 도메인 추출 실패 시 기본값
        except Exception as e:
            self.logger.warning(f"EPUB 저자 설정 중 도메인 추출 오류: {e}. 기본 저자 사용.")
            book.add_author("Unknown")

        # CSS 스타일 추가
        style = f'''
        body {{ 
            font-family: "{self.font_family}", sans-serif;
            line-height: 1.8;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 2em;
        }}
        img, video {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 1.5em auto;
            border-radius: 4px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        pre, code {{
            max-width: 100%;
            overflow-x: auto;
            word-break: break-all;
            white-space: pre-wrap;
            box-sizing: border-box;
        }}
        h1 {{ 
            font-size: 2em;
            margin: 1.5em 0 1em;
            color: #1a1a1a;
            border-bottom: 1px solid #eaecef;
            padding-bottom: 0.3em;
        }}
        h2 {{ 
            font-size: 1.5em;
            margin: 1.2em 0 0.8em;
            color: #1a1a1a;
        }}
        h3 {{ 
            font-size: 1.25em;
            margin: 1em 0 0.6em;
            color: #1a1a1a;
        }}
        p {{
            margin: 0.8em 0;
            line-height: 1.8;
        }}
        pre {{
            background-color: #f6f8fa;
            border-radius: 6px;
            padding: 16px;
            overflow: auto;
            font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, monospace;
            font-size: 0.9em;
            line-height: 1.45;
            margin: 1em 0;
            border: 1px solid #ddd; /* 코드 블록 테두리 추가 */
            border-radius: 4px; /* 코드 블록 모서리 둥글게 */
        }}
        code {{ /* 인라인 코드 스타일 */
            font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, monospace;
            background-color: #e7e7e7; /* 인라인 코드 배경색 약간 변경 */
            padding: 0.2em 0.4em;
            border-radius: 3px;
            font-size: 0.9em;
            word-break: break-all; /* 인라인 코드도 긴 경우 줄바꿈 */
        }}
        pre code {{ /* 코드 블록 내의 code 태그는 pre 스타일을 따르도록 */
            background-color: transparent;
            padding: 0;
            border-radius: 0;
            border: none; 
            font-size: inherit; /* pre의 폰트 크기 상속 */
            line-height: inherit; /* pre의 줄 간격 상속 */
        }}
        blockquote {{
            margin: 1em 0;
            padding: 0 1em;
            color: #6a737d;
            border-left: 0.25em solid #dfe2e5;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
            table-layout: fixed; /* 표 레이아웃 고정 */
            border: 1px solid #ccc; /* 표 전체 테두리 */
        }}
        th, td {{
            border: 1px solid #ccc; /* 셀 테두리 변경 */
            padding: 8px 10px; /* 셀 패딩 조정 */
            word-wrap: break-word; /* 긴 단어 줄바꿈 (overflow-wrap과 유사) */
            overflow-wrap: break-word; /* 내용이 셀을 넘칠 때 줄바꿈 */
            hyphens: auto; /* 필요한 경우 하이픈 추가 (지원하는 뷰어에서) */
            text-align: left; /* 기본 왼쪽 정렬 */
        }}
        th {{
            background-color: #f6f8fa;
        }}
        hr {{
            height: 0.25em;
            padding: 0;
            margin: 24px 0;
            background-color: #e1e4e8;
            border: 0;
        }}
        '''
        css = epub.EpubItem(uid="style_default", file_name="style/default.css", media_type="text/css", content=style)
        book.add_item(css)
        
        # 표지 페이지 생성
        cover = epub.EpubHtml(title='표지', file_name='cover.xhtml', lang='ko')
        cover.content = f'''
        <html>
        <head>
            <link rel="stylesheet" href="style/default.css" type="text/css" />
        </head>
        <body>
            <h1>{self.title}</h1>
        </body>
        </html>
        '''
        book.add_item(cover)
        
        # 목차 페이지 생성
        if len(self.pages) > 1:
            index = epub.EpubHtml(title='목차', file_name='index.xhtml', lang='ko')
            index.content = f'''
            <html>
            <head>
                <link rel="stylesheet" href="style/default.css" type="text/css" />
            </head>
            <body>
                {self.create_index_page()}
            </body>
            </html>
            '''
            book.add_item(index)
        
        # 각 페이지 처리
        chapters = []
        for i, page in enumerate(self.pages, 1):
            chapter = epub.EpubHtml(
                title=page.get('title', f'페이지 {i}'),
                file_name=f'page{i}.xhtml',
                lang='ko'
            )
            
            # 이미지 처리
            soup = BeautifulSoup(self.clean_for_reading_mode(page['content']), 'html.parser')
            for img in soup.find_all('img'):
                img_url = urljoin(self.base_url, img.get('src', ''))
                try:
                    response = requests.get(img_url)
                    if response.status_code == 200:
                        img_filename = f"images/{hash(img_url)}.png"
                        img_item = epub.EpubItem(
                            uid=f"img_{hash(img_url)}",
                            file_name=img_filename,
                            media_type="image/png",
                            content=response.content
                        )
                        book.add_item(img_item)
                        img['src'] = img_filename
                        
                        # 이미지 캡션 추가
                        if img.get('alt'):
                            caption = soup.new_tag('p')
                            caption['class'] = 'image-caption'
                            caption.string = img['alt']
                            img.insert_after(caption)
                except Exception as e:
                    self.logger.error(f"이미지 처리 중 오류 발생: {str(e)}")
            
            # 코드 블록 처리
            for pre in soup.find_all('pre'):
                if pre.find('code'):
                    code = pre.find('code')
                    # 언어 지정이 있는 경우
                    if code.get('class'):
                        lang = code['class'][0].replace('language-', '')
                        pre['data-lang'] = lang
            
            chapter.content = f'''
            <html>
            <head>
                <link rel="stylesheet" href="style/default.css" type="text/css" />
            </head>
            <body>
                {str(soup)}
            </body>
            </html>
            '''
            book.add_item(chapter)
            chapters.append(chapter)
        
        # 목차 생성
        book.toc = [cover]
        if len(self.pages) > 1:
            book.toc.append(index)
        book.toc.extend(chapters)
        
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # 스파인 정의
        book.spine = ['nav', cover]
        if len(self.pages) > 1:
            book.spine.append(index)
        book.spine.extend(chapters)
        
        # EPUB 파일 생성
        epub.write_epub(output_path, book, {})

    def create_markdown(self, content: str, output_path: str):
        """Markdown 형식의 문서를 생성합니다."""
        if len(self.pages) > 1:
            # ZIP 파일로 묶기
            with zipfile.ZipFile(output_path, 'w') as zipf:
                # 목차 생성
                index_content = "# 목차\n\n"
                for i, page in enumerate(self.pages, 1):
                    title = page.get('title', f'페이지 {i}')
                    index_content += f"{i}. [{title}]({i:02d}_{title}.md)\n"
                
                # 목차 파일 추가
                zipf.writestr("00_index.md", index_content)
                
                # 각 페이지 처리
                for i, page in enumerate(self.pages, 1):
                    title = page.get('title', f'페이지 {i}')
                    filename = f"{i:02d}_{title}.md"
                    
                    soup = BeautifulSoup(self.clean_for_reading_mode(page['content']), 'html.parser')
                    md_content = ""
                    
                    # 모든 요소를 순서대로 처리
                    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'pre', 'blockquote', 'img', 'table']):
                        if element.name.startswith('h'):
                            level = int(element.name[1])
                            md_content += f"{'#' * level} {element.get_text().strip()}\n\n"
                        elif element.name == 'p':
                            text = element.get_text().strip()
                            if text:
                                md_content += f"{text}\n\n"
                        elif element.name == 'pre':
                            code = element.find('code')
                            if code:
                                lang = code.get('class', [''])[0].replace('language-', '') if code.get('class') else ''
                                md_content += f"```{lang}\n{code.get_text()}\n```\n\n"
                            else:
                                md_content += f"```\n{element.get_text()}\n```\n\n"
                        elif element.name == 'blockquote':
                            md_content += f"> {element.get_text().strip()}\n\n"
                        elif element.name == 'img':
                            alt = element.get('alt', '')
                            src = urljoin(self.base_url, element.get('src', ''))
                            md_content += f"![{alt}]({src})\n\n"
                            if alt:
                                md_content += f"*{alt}*\n\n"
                        elif element.name == 'table':
                            md_content += self._convert_table_to_markdown(element)
                    
                    zipf.writestr(filename, md_content)
        else:
            # 단일 페이지 처리
            if not output_path.lower().endswith(".md"):
                output_path += ".md"
            soup = BeautifulSoup(self.clean_for_reading_mode(content), 'html.parser')
            with open(output_path, 'w', encoding='utf-8') as f:
                for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'pre', 'blockquote', 'img', 'table']):
                    if element.name.startswith('h'):
                        level = int(element.name[1])
                        f.write(f"{'#' * level} {element.get_text().strip()}\n\n")
                    elif element.name == 'p':
                        text = element.get_text().strip()
                        if text:
                            f.write(f"{text}\n\n")
                    elif element.name == 'pre':
                        code = element.find('code')
                        if code:
                            lang = code.get('class', [''])[0].replace('language-', '') if code.get('class') else ''
                            f.write(f"```{lang}\n{code.get_text()}\n```\n\n")
                        else:
                            f.write(f"```\n{element.get_text()}\n```\n\n")
                    elif element.name == 'blockquote':
                        f.write(f"> {element.get_text().strip()}\n\n")
                    elif element.name == 'img':
                        alt = element.get('alt', '')
                        src = urljoin(self.base_url, element.get('src', ''))
                        f.write(f"![{alt}]({src})\n\n")
                        if alt:
                            f.write(f"*{alt}*\n\n")
                    elif element.name == 'table':
                        f.write(self._convert_table_to_markdown(element))

    def _convert_table_to_markdown(self, table):
        """HTML 테이블을 Markdown 형식으로 변환합니다."""
        md = []
        rows = table.find_all('tr')
        
        # 헤더 처리
        headers = []
        for th in rows[0].find_all(['th', 'td']):
            headers.append(th.get_text().strip())
        md.append('| ' + ' | '.join(headers) + ' |')
        md.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
        
        # 데이터 행 처리
        for row in rows[1:]:
            cells = []
            for cell in row.find_all(['td', 'th']):
                cells.append(cell.get_text().strip())
            md.append('| ' + ' | '.join(cells) + ' |')
        
        return '\n'.join(md) + '\n\n'

    def create_doc(self, content: str, output_path: str):
        """Word 문서 형식으로 저장합니다."""
        doc = Document()
        
        # 기본 스타일 설정
        styles = doc.styles
        styles['Normal'].font.name = self.font_family
        styles['Normal'].font.size = Pt(11)
        styles['Normal'].paragraph_format.line_spacing = 1.5
        
        # 제목 스타일 설정
        for i in range(1, 7):
            style_name = f'Heading {i}'
            if style_name in styles:
                styles[style_name].font.name = self.font_family
                styles[style_name].font.size = Pt(16 - i)
                styles[style_name].font.bold = True
        
        # 코드 블록 스타일 설정
        if 'Code' not in styles:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = 'Consolas'
            code_style.font.size = Pt(10)
            code_style.paragraph_format.space_before = Pt(12)
            code_style.paragraph_format.space_after = Pt(12)
            code_style.paragraph_format.left_indent = Pt(24)
            code_style.paragraph_format.right_indent = Pt(24)
            code_style.paragraph_format.background_color = RGBColor(246, 248, 250)
        
        soup = BeautifulSoup(self.clean_for_reading_mode(content), 'html.parser')
        
        # 모든 요소를 순서대로 처리
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'pre', 'blockquote', 'img', 'table']):
            if element.name.startswith('h'):
                # 제목 태그 처리
                level = int(element.name[1])
                doc.add_heading(element.get_text().strip(), level=level)
            elif element.name == 'p':
                # 단락 처리
                text = element.get_text().strip()
                if text:  # 빈 단락은 건너뛰기
                    p = doc.add_paragraph(text)
                    p.style = 'Normal'
            elif element.name == 'pre':
                # 코드 블록 처리
                code = element.find('code')
                if code:
                    lang = code.get('class', [''])[0].replace('language-', '') if code.get('class') else ''
                    p = doc.add_paragraph(code.get_text(), style='Code')
                    if lang:
                        p.add_run(f'\n// {lang}').italic = True
                else:
                    doc.add_paragraph(element.get_text(), style='Code')
            elif element.name == 'blockquote':
                # 인용구 처리
                p = doc.add_paragraph(element.get_text().strip())
                p.style = 'Quote'
            elif element.name == 'img':
            # 이미지 처리
                try:
                    img_url = urljoin(self.base_url, element.get('src', ''))
                    alt_text = element.get('alt', '')

                    if img_url.lower().endswith('.svg'):
                        self.logger.warning(f"SVG 이미지는 DOCX에서 직접 지원되지 않습니다. 건너뜁니다: {img_url}")
                        caption_text = f"[SVG 이미지: {alt_text or img_url}]"
                        if alt_text:
                            p = doc.add_paragraph(caption_text)
                            p.style = 'Caption'
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        continue

                    response = requests.get(img_url)
                    response.raise_for_status() # HTTP 오류 발생 시 예외 발생

                    # 이미지 형식을 유추하거나 기본값 사용 (PNG, JPG 등)
                    # content_type = response.headers.get('Content-Type', '').lower()
                    # if 'jpeg' in content_type or 'jpg' in content_type:
                    #     suffix = '.jpg'
                    # elif 'png' in content_type:
                    #     suffix = '.png'
                    # else: # 기타 지원될 수 있는 형식 또는 알 수 없는 형식
                    #     self.logger.warning(f"알 수 없는 이미지 Content-Type ({content_type}) 또는 지원되지 않는 이미지 형식일 수 있습니다: {img_url}. PNG로 가정하고 시도합니다.")
                    #     suffix = '.png' # 기본적으로 png로 시도

                    # 파일 확장자로 이미지 형식 결정 시도
                    img_filename_lower = img_url.lower()
                    if img_filename_lower.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff')):
                        suffix = os.path.splitext(img_url)[1]
                    else:
                        # Content-Type 확인 시도 (선택적)
                        content_type = response.headers.get('Content-Type', '').lower()
                        if 'jpeg' in content_type or 'jpg' in content_type:
                            suffix = '.jpg'
                        elif 'png' in content_type:
                            suffix = '.png'
                        else:
                            self.logger.warning(f"이미지 URL에 명확한 확장자가 없고 Content-Type({content_type})으로도 유추하기 어렵습니다. PNG로 가정: {img_url}")
                            suffix = '.png'


                    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                        tmp.write(response.content)
                        tmp_path = tmp.name
                    
                    try:
                        # 이미지 너비는 페이지 너비에 맞게 조정 (예: 6인치)
                        doc.add_picture(tmp_path, width=Inches(6))
                        if alt_text:
                            p = doc.add_paragraph(alt_text)
                            p.style = 'Caption'
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e_img_add:
                        self.logger.error(f"DOCX에 이미지 추가 중 오류 발생 (파일: {tmp_path}, URL: {img_url}): {repr(e_img_add)}")
                        doc.add_paragraph(f"[이미지 삽입 실패: {alt_text or img_url} - {repr(e_img_add)}]")
                    finally:
                        if os.path.exists(tmp_path):
                            os.remove(tmp_path)

                except requests.exceptions.RequestException as e_req:
                    self.logger.error(f"DOCX 이미지 다운로드 중 오류 발생 (URL: {img_url}): {repr(e_req)}")
                    doc.add_paragraph(f"[이미지 다운로드 실패: {alt_text or img_url} - {repr(e_req)}]")
                except Exception as e_generic:
                    self.logger.error(f"DOCX 이미지 처리 중 알 수 없는 오류 발생 (URL: {img_url}): {repr(e_generic)}")
                    doc.add_paragraph(f"[이미지 처리 중 알 수 없는 오류: {alt_text or img_url} - {repr(e_generic)}]")
            elif element.name == 'table':
                # 테이블 처리
                table_data = []
                for tr in element.find_all('tr'):
                    row = []
                    for td in tr.find_all(['td', 'th']): # th도 데이터로 처리할 수 있도록 수정
                        row.append(td.get_text().strip())
                    table_data.append(row)
                
                if table_data:
                    # 테이블 생성 (첫 행을 헤더로 가정)
                    num_cols = len(table_data[0]) if table_data else 0
                    if num_cols > 0:
                        # 행 수를 실제 데이터에 맞게 조정
                        num_rows = len(table_data)
                        try:
                            doc_table = doc.add_table(rows=num_rows, cols=num_cols)
                            doc_table.style = 'TableGrid' # 기본 표 스타일 적용
                            doc_table.autofit = True # 자동 맞춤 기능 (내용에 따라 너비 조절 시도)

                            for i, row_data in enumerate(table_data):
                                cells = doc_table.rows[i].cells
                                for j, cell_text in enumerate(row_data):
                                    if j < len(cells): # 열 인덱스 범위 확인
                                        cells[j].text = cell_text
                                    else:
                                        self.logger.warning(f"테이블 데이터 열 ({j+1})이 실제 테이블 열 수 ({len(cells)})를 초과합니다.")
                            
                            # 첫 행을 헤더로 처리 (선택적: 굵게 등)
                            if num_rows > 0:
                                for cell in doc_table.rows[0].cells:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
                        except Exception as e_table:
                            self.logger.error(f"DOCX 테이블 생성 중 오류: {e_table}")
                            doc.add_paragraph(f"[테이블 생성 오류: {e_table}]")
                    else:
                        self.logger.warning("테이블 데이터가 비어있거나 형식이 잘못되어 DOCX 테이블을 생성할 수 없습니다.")
                else:
                    self.logger.info("HTML에 테이블 데이터가 없어 DOCX 테이블을 생성하지 않습니다.")
        
        doc.save(output_path)

    def create_pdf(self, content: str, output_path: str):
        """PDF 형식의 문서를 생성합니다."""
        soup = BeautifulSoup(self.clean_for_reading_mode(content), 'html.parser')
        
        # 이미지 처리
        temp_image_files = []
        for img in soup.find_all('img'):
            try:
                img_url = urljoin(self.base_url, img.get('src', ''))
                if img_url.startswith('file://'): # Already a local file, likely from a previous step or bad source
                    pass
                else:
                    response = requests.get(img_url, timeout=10) # Added timeout
                    if response.status_code == 200:
                        # 임시 파일로 이미지 저장
                        tmp_img_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                        tmp_img_file.write(response.content)
                        tmp_img_file.close()
                        # 이미지 태그의 src를 로컬 파일 경로로 변경
                        img['src'] = f"file://{tmp_img_file.name}"
                        temp_image_files.append(tmp_img_file.name)
                    else:
                        self.logger.warning(f"이미지 다운로드 실패 ({response.status_code}): {img_url}")
                        img.replace_with(soup.new_string(f"[이미지 다운로드 실패: {img.get('alt', '')}]"))
            except Exception as e:
                self.logger.error(f"PDF 이미지 처리 중 오류 발생: {str(e)}")
                # 이미지 로드 실패 시 alt 텍스트로 대체
                img.replace_with(soup.new_string(f"[이미지 처리 오류: {img.get('alt', '')}]"))
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@400;700&display=swap');
                @page {{
                    margin: 2.5cm;
                    @top-right {{
                        content: counter(page);
                    }}
                }}
                body {{ 
                    font-family: '{self.font_family}', sans-serif;
                    line-height: 1.8;
                    color: #333;
                    max-width: 800px;
                    margin: 0 auto;
                }}
                h1 {{ 
                    font-size: 2em;
                    margin: 1.5em 0 1em;
                    color: #1a1a1a;
                    border-bottom: 1px solid #eaecef;
                    padding-bottom: 0.3em;
                }}
                h2 {{ 
                    font-size: 1.5em;
                    margin: 1.2em 0 0.8em;
                    color: #1a1a1a;
                }}
                h3 {{ 
                    font-size: 1.25em;
                    margin: 1em 0 0.6em;
                    color: #1a1a1a;
                }}
                p {{
                    margin: 0.8em 0;
                    line-height: 1.8;
                }}
                pre {{
                    background-color: #f6f8fa;
                    border-radius: 6px;
                    padding: 16px;
                    overflow: auto;
                    font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, monospace;
                    font-size: 0.9em;
                    line-height: 1.45;
                    margin: 1em 0;
                    white-space: pre-wrap; /* 코드 줄바꿈 허용 */
                    word-break: break-word; /* 긴 코드 줄바꿈 처리 */
                }}
                code {{
                    font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, monospace;
                    background-color: rgba(27, 31, 35, 0.05);
                    padding: 0.2em 0.4em;
                    border-radius: 3px;
                    font-size: 0.9em;
                }}
                pre code {{
                    background-color: transparent;
                    padding: 0;
                    border-radius: 0;
                }}
                img {{
                    max-width: 100%;
                    height: auto;
                    display: block;
                    margin: 1.5em auto;
                    border-radius: 4px;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                }}
                .image-caption {{
                    text-align: center;
                    color: #666;
                    font-style: italic;
                    margin-top: 0.5em;
                }}
                blockquote {{
                    margin: 1em 0;
                    padding: 0 1em;
                    color: #6a737d;
                    border-left: 0.25em solid #dfe2e5;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 1em 0;
                }}
                th, td {{
                    border: 1px solid #dfe2e5;
                    padding: 6px 13px;
                    word-break: break-word; /* 긴 단어 줄바꿈 추가 */
                }}
                th {{
                    background-color: #f6f8fa;
                }}
                hr {{
                    height: 0.25em;
                    padding: 0;
                    margin: 24px 0;
                    background-color: #e1e4e8;
                    border: 0;
                }}
            </style>
        </head>
        <body>
            {str(soup)}
        </body>
        </html>
        """
        
        try:
            # PDF 생성
            HTML(string=html_content, base_url=self.base_url).write_pdf(
                output_path,
                presentational_hints=True,
                optimize_size=('fonts', 'images'),
                zoom=1
            )
            self.logger.info(f"PDF 파일이 생성되었습니다: {output_path}")
        except Exception as e:
            self.logger.error(f"PDF 생성 중 오류 발생: {str(e)}")
            raise
        finally:
            # 임시 이미지 파일 삭제
            for tmp_file_path in temp_image_files:
                if os.path.exists(tmp_file_path):
                    try:
                        os.remove(tmp_file_path)
                        self.logger.debug(f"임시 이미지 파일 삭제: {tmp_file_path}")
                    except Exception as e_remove:
                        self.logger.error(f"임시 이미지 파일 삭제 중 오류: {tmp_file_path}, {e_remove}")

    def translate_text(self, text: str, target_lang: str) -> str:
        """텍스트를 지정된 언어로 번역합니다."""
        try:
            translator = GoogleTranslator(source='auto', target=target_lang)
            return translator.translate(text)
        except Exception as e:
            self.logger.error(f"번역 중 오류 발생: {str(e)}")
            return text

    def translate_html_text(self, html, translator_type): # target_lang 파라미터 제거
        # 'browser' 또는 'chrome_builtin' (이전 버전 호환)의 경우, 이미 Selenium 단계에서 처리되었으므로 추가 작업 없음
        if translator_type == 'browser' or translator_type == 'chrome_builtin':
            self.logger.info(f"'{translator_type}' 옵션 선택됨. Selenium 단계에서 브라우저 번역이 시도되었으므로 서버 측 추가 번역은 수행하지 않습니다.")
            return html

        if not self.target_lang:
            self.logger.warning("대상 언어가 설정되지 않아 서버 측 번역을 건너<0xEB><0x9B><0x8D>니다.")
            return html

        soup = BeautifulSoup(html, 'html.parser')
        
        for elem in soup.find_all(text=True):
            parent = elem.parent.name
            if parent not in ['code', 'pre', 'script', 'style', 'table', 'a']:
                text_to_translate = elem.strip()
                if text_to_translate:
                    try:
                        translated_text = None
                        if translator_type == 'papago':
                            if self.papago_id and self.papago_secret:
                                translated_text = self.translate_with_papago(text_to_translate, self.target_lang)
                            else:
                                self.logger.warning("Papago API 키가 없어 Papago 번역을 건너<0xEB><0x9B><0x8D>니다.")
                                translated_text = text_to_translate
                        elif translator_type == 'gpt':
                            if self.openai_key:
                                translated_text = self.translate_with_gpt(text_to_translate, self.target_lang)
                            else:
                                self.logger.warning("OpenAI API 키가 없어 GPT 번역을 건너<0xEB><0x9B><0x8D>니다.")
                                translated_text = text_to_translate
                        elif translator_type == 'deepl':
                            if self.deepl_key:
                                translated_text = self.translate_with_deepl(text_to_translate, self.target_lang)
                            else:
                                self.logger.warning("DeepL API 키가 없어 DeepL 번역을 건너<0xEB><0x9B><0x8D>니다.")
                                translated_text = text_to_translate
                        elif translator_type == 'google': # 'google' 명시
                            self.logger.debug(f"GoogleTranslator 호출: 원본='{text_to_translate[:30]}...', 대상='{self.target_lang}'")
                            try:
                                translated_candidate = GoogleTranslator(source='auto', target=self.target_lang).translate(text_to_translate)
                                if translated_candidate and translated_candidate.strip():
                                    if translated_candidate.strip().lower() != text_to_translate.strip().lower():
                                        translated_text = translated_candidate
                                    else:
                                        translated_text = text_to_translate
                                else:
                                    translated_text = text_to_translate
                            except Exception as e_google:
                                self.logger.error(f"GoogleTranslator 오류: {str(e_google)}. 원본 유지.")
                                translated_text = text_to_translate
                        else: # 알 수 없는 번역기 또는 'none' (이미 process에서 처리됨)
                            translated_text = text_to_translate

                        if translated_text is not None and translated_text.strip() != "":
                           elem.replace_with(translated_text)
                        # else: # 번역 결과가 비거나 원본과 같으면 원본 유지 (별도 처리 불필요)

                    except Exception as e:
                        self.logger.error(f"번역 중 오류 (번역기: {translator_type}, 대상: {self.target_lang}): {str(e)}. 원본 유지.")
        
        return str(soup)

    def translate_with_papago(self, text, target_lang):
        url = "https://openapi.naver.com/v1/papago/n2mt"
        headers = {
            "X-Naver-Client-Id": self.papago_id,
            "X-Naver-Client-Secret": self.papago_secret
        }
        data = {
            "source": "ko" if target_lang != "ko" else "en",
            "target": target_lang,
            "text": text
        }
        response = requests.post(url, headers=headers, data=data)
        if response.status_code == 200:
            return response.json()['message']['result']['translatedText']
        else:
            return text

    def translate_with_gpt(self, text, target_lang):
        import openai
        openai.api_key = self.openai_key
        prompt = f"Translate the following text to {target_lang}:\n{text}"
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}]
            )
            return response.choices[0].message['content'].strip()
        except Exception as e:
            self.logger.error(f'GPT 번역 오류: {e}')
            return text

    def translate_with_deepl(self, text, target_lang):
        url = "https://api-free.deepl.com/v2/translate"
        headers = {"Authorization": f"DeepL-Auth-Key {self.deepl_key}"}
        data = {"text": text, "target_lang": target_lang.upper()}
        response = requests.post(url, headers=headers, data=data)
        if response.status_code == 200:
            return response.json()['translations'][0]['text']
        else:
            return text

    def clean_for_reading_mode(self, html):
        soup = BeautifulSoup(html, 'html.parser')

        # SVG, 아이콘, 불필요한 장식 제거 및 대체
        for svg in soup.find_all('svg'):
            svg.decompose()
        for icon in soup.find_all(['i', 'span']):
            if 'icon' in icon.get('class', []) or 'material-icons' in icon.get('class', []):
                alt = icon.get('aria-label') or icon.get('title') or '[ICON]'
                icon.replace_with(soup.new_string(alt))
        # 불필요한 nav, aside, footer, button, form, script, style, header, noscript, iframe, input, select, textarea, label, ul.menu, div.menu 등 제거
        for tag in soup.find_all(['nav', 'aside', 'footer', 'button', 'form', 'script', 'style', 'header', 'noscript', 'iframe', 'input', 'select', 'textarea', 'label']):
            tag.decompose()
        for tag in soup.find_all(class_=['menu', 'sidebar', 'navigation', 'breadcrumb', 'toc', 'toolbar', 'pagination', 'logo', 'search', 'ads', 'ad', 'banner', 'sponsor']):
            tag.decompose()
        # 코드블록 언어 표시
        for pre in soup.find_all('pre'):
            code = pre.find('code')
            if code and code.get('class'):
                lang = code['class'][0].replace('language-', '')
                pre['data-lang'] = lang
        # 이미지 alt 캡션 추가
        for img in soup.find_all('img'):
            if img.get('alt'):
                caption = soup.new_tag('p')
                caption['class'] = 'image-caption'
                caption.string = img['alt']
                img.insert_after(caption)
        # 본문, 제목, 표, 코드, 이미지, 인용구만 남기고 나머지 제거
        allowed = ['h1','h2','h3','h4','h5','h6','p','pre','code','img','blockquote','ul','ol','li','table','thead','tbody','tr','th','td','hr']
        for tag in soup.find_all(True):
            if tag.name not in allowed:
                tag.unwrap()
        # [SVG] 문자열 후처리로 모두 삭제
        result = str(soup)
        result = re.sub(r'\s*\[SVG\]\s*', '', result, flags=re.IGNORECASE)
        return result

    def process(self, output_path: str, output_format: str = 'epub'): # target_lang 파라미터 제거
        """전체 프로세스를 실행합니다."""
        try:
            # URL에서 웹페이지 타이틀 가져오기 (만약 self.title이 비어있다면)
            if not self.title and self.base_url:
                try:
                    self.driver.get(self.base_url)
                    time.sleep(1) # 페이지 로딩 대기
                    page_title = self.driver.title
                    if page_title:
                        self.title = page_title.strip()
                    else:
                        # 타이틀이 없는 경우를 대비한 기본값
                        self.title = "제목 없음"
                    self.logger.info(f"웹페이지에서 추출한 문서 제목: {self.title}")
                except Exception as e:
                    self.logger.warning(f"웹페이지 제목 추출 중 오류 발생: {e}. 기본 제목을 사용합니다.")
                    if not self.title: # 오류 발생 시에도 제목이 비어있으면 기본값 설정
                        self.title = "제목 없음"
            elif not self.title and not self.base_url: # URL도 없고 텍스트도 없는 경우 (일어나면 안되지만 방어 코드)
                 self.title = "제목 없음"

            self.logger.info("페이지 내용 수집 중...")
            
            # 메뉴 선택자가 있고, content_selector가 AUTO가 아닐 때 (즉, 여러 페이지 모드일 때)만 메뉴 링크 처리
            if self.menu_selector and self.content_selector != "__AUTO_SINGLE_PAGE__":
                menu_links = self.get_menu_links()
                for link in menu_links:
                    page_content = self.get_page_content(link)
                    if page_content:
                        self.pages.append({
                            'title': link.split('/')[-1],
                            'content': page_content
                        })
            else:
                # 메뉴 선택자가 없는 경우 메인 페이지만 처리
                main_content = self.get_page_content()
                self.pages.append({
                    'title': self.title,
                    'content': main_content
                })
            
            # 번역 처리
            if self.target_lang and self.pages: # self.target_lang 사용
                # 'browser' 또는 'chrome_builtin'이 아닌 경우에만 서버 측 번역 시도
                if self.translator_type not in ['none', 'browser', 'chrome_builtin']:
                    for i, page in enumerate(self.pages):
                        self.logger.info(f"페이지 {i+1} 서버 측 번역 시도. 번역기: {self.translator_type}, 대상 언어: {self.target_lang}")
                        page['content'] = self.translate_html_text(page['content'], self.translator_type)
                elif self.translator_type in ['browser', 'chrome_builtin']:
                    self.logger.info(f"'{self.translator_type}' 옵션이 선택되어 Selenium 단계에서 브라우저 번역이 시도되었습니다. 서버 측 추가 번역은 수행하지 않습니다.")
                else: # self.translator_type == 'none'
                    self.logger.info("번역 옵션이 'none'으로 설정되어 번역을 수행하지 않습니다.")
            
            self.logger.info(f"{output_format.upper()} 파일 생성 중...")
            
            # 출력 형식에 따라 적절한 변환 함수 호출
            if output_format.lower() == 'epub':
                self.create_epub(self.pages[0]['content'], output_path)
            elif output_format.lower() == 'markdown':
                self.create_markdown(self.pages[0]['content'], output_path)
            elif output_format.lower() == 'doc':
                self.create_doc(self.pages[0]['content'], output_path)
            elif output_format.lower() == 'pdf':
                self.create_pdf(self.pages[0]['content'], output_path)
            else:
                raise ValueError(f"지원하지 않는 출력 형식입니다: {output_format}")
            
            self.logger.info(f"문서가 생성되었습니다: {output_path}")
            
        except Exception as e:
            self.logger.error(f"처리 중 오류 발생: {str(e)}")
        finally:
            self.driver.quit()

def main():
    # 사용 예시
    converter = WebToEbook(
        title="Cursor Documentation",
        base_url="https://docs.cursor.com/",
        content_selector="main",
        menu_selector="nav",  # 메뉴 선택자 추가
        translator_type="chrome_builtin", # 예: 크롬 내장 번역 사용
        target_lang="en",                 # 대상 언어 설정
        papago_id=None,
        papago_secret=None,
        openai_key=None,
        deepl_key=None,
        font_family='Noto Sans KR'
    )
    
    converter.process(
        output_path="cursor_docs.pdf", # target_lang 인자 제거
        output_format="pdf"
    )

if __name__ == "__main__":
    main()
